"""
Main entry point for the Word Document MCP Server.
Acts as the central controller for the MCP server that handles Word document operations.
Supports multiple transports: stdio, sse, and streamable-http using standalone FastMCP.
"""

import os
import sys
# Set required environment variable for FastMCP 2.8.1+
os.environ.setdefault('FASTMCP_LOG_LEVEL', 'INFO')
from fastmcp import FastMCP
from word_document_server.tools import (
    document_tools,
    content_tools,
    format_tools,
    protection_tools,
    footnote_tools,
    extended_document_tools
)
from word_document_server.tools import batch_content_tools
from word_document_server.utils.file_utils import download_file_from_url, upload_file_to_server
from typing import Optional

def get_transport_config():
    """
    Get transport configuration from environment variables.
    
    Returns:
        dict: Transport configuration with type, host, port, and other settings
    """
    # Default configuration
    config = {
        'transport': 'stdio',  # Default to stdio for backward compatibility
        'host': '127.0.0.1',
        'port': 8000,
        'path': '/mcp',
        'sse_path': '/sse'
    }
    
    # Override with environment variables if provided
    transport = os.getenv('MCP_TRANSPORT', 'stdio').lower()
    print(f"Transport: {transport}")
    # Validate transport type
    valid_transports = ['stdio', 'streamable-http', 'sse']
    if transport not in valid_transports:
        print(f"Warning: Invalid transport '{transport}'. Falling back to 'stdio'.")
        transport = 'stdio'
    
    config['transport'] = transport
    config['host'] = os.getenv('MCP_HOST', config['host'])
    config['port'] = int(os.getenv('MCP_PORT', config['port']))
    config['path'] = os.getenv('MCP_PATH', config['path'])
    config['sse_path'] = os.getenv('MCP_SSE_PATH', config['sse_path'])
    
    return config


def setup_logging(debug_mode):
    """
    Setup logging based on debug mode.
    
    Args:
        debug_mode (bool): Whether to enable debug logging
    """
    import logging
    
    if debug_mode:
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        print("Debug logging enabled")
    else:
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )


# Initialize FastMCP server
mcp = FastMCP("Word Document Server")


def register_tools():
    """Register all tools with the MCP server using FastMCP decorators."""
    
    # Document tools (create, copy, info, etc.)
    @mcp.tool()
    def create_document(filename: str, title: str = None, author: str = None):
        """Create a new Word document with optional metadata."""
        return document_tools.create_document(filename, title, author)
    
    @mcp.tool()
    def copy_document(source_filename: str, destination_filename: str = None):
        """Create a copy of a Word document."""
        return document_tools.copy_document(source_filename, destination_filename)
    
    @mcp.tool()
    def get_document_info(filename: str):
        """Get information about a Word document."""
        return document_tools.get_document_info(filename)
    
    @mcp.tool()
    def get_document_text(filename: str):
        """Extract all text from a Word document."""
        return document_tools.get_document_text(filename)
    
    @mcp.tool()
    def get_document_outline(filename: str):
        """Get the structure of a Word document."""
        return document_tools.get_document_outline(filename)
    
    @mcp.tool()
    def list_available_documents(directory: str = "."):
        """List all .docx files in the specified directory."""
        return document_tools.list_available_documents(directory)
    
    # Content tools (paragraphs, headings, tables, etc.)
    @mcp.tool()
    def add_paragraph(filename: str, text: str, style: str = None):
        """Add a paragraph to a Word document."""
        return content_tools.add_paragraph(filename, text, style)
    
    @mcp.tool()
    def add_heading(filename: str, text: str, level: int = 1):
        """Add a heading to a Word document."""
        return content_tools.add_heading(filename, text, level)
    
    @mcp.tool()
    def add_picture(filename: str, image_path: str, width: float = None):
        """Add an image to a Word document."""
        return content_tools.add_picture(filename, image_path, width)
    
    @mcp.tool()
    def add_table(filename: str, rows: int, cols: int, data: list = None):
        """Add a table to a Word document."""
        return content_tools.add_table(filename, rows, cols, data)
    
    @mcp.tool()
    def add_page_break(filename: str):
        """Add a page break to the document."""
        return content_tools.add_page_break(filename)
    
    @mcp.tool()
    def delete_paragraph(filename: str, paragraph_index: int):
        """Delete a paragraph from a document."""
        return content_tools.delete_paragraph(filename, paragraph_index)
    
    @mcp.tool()
    def search_and_replace(filename: str, find_text: str, replace_text: str):
        """Search for text and replace all occurrences."""
        return content_tools.search_and_replace(filename, find_text, replace_text)
    
    # Format tools (styling, text formatting, etc.)
    @mcp.tool()
    def create_custom_style(filename: str, style_name: str, bold: bool = None, 
                          italic: bool = None, font_size: int = None, 
                          font_name: str = None, color: str = None, 
                          base_style: str = None):
        """Create a custom style in the document."""
        return format_tools.create_custom_style(
            filename, style_name, bold, italic, font_size, font_name, color, base_style
        )
    
    @mcp.tool()
    def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int,
                   bold: bool = None, italic: bool = None, underline: bool = None,
                   color: str = None, font_size: int = None, font_name: str = None):
        """Format a specific range of text within a paragraph."""
        return format_tools.format_text(
            filename, paragraph_index, start_pos, end_pos, bold, italic, 
            underline, color, font_size, font_name
        )
    
    @mcp.tool()
    def format_table(filename: str, table_index: int, has_header_row: bool = None,
                    border_style: str = None, shading: list = None):
        """Format a table with borders, shading, and structure."""
        return format_tools.format_table(filename, table_index, has_header_row, border_style, shading)
    
    # Protection tools
    @mcp.tool()
    def protect_document(filename: str, password: str):
        """Add password protection to a Word document."""
        return protection_tools.protect_document(filename, password)
    
    @mcp.tool()
    def unprotect_document(filename: str, password: str):
        """Remove password protection from a Word document."""
        return protection_tools.unprotect_document(filename, password)
    
    # Footnote tools
    @mcp.tool()
    def add_footnote_to_document(filename: str, paragraph_index: int, footnote_text: str):
        """Add a footnote to a specific paragraph in a Word document."""
        return footnote_tools.add_footnote_to_document(filename, paragraph_index, footnote_text)
    
    @mcp.tool()
    def add_endnote_to_document(filename: str, paragraph_index: int, endnote_text: str):
        """Add an endnote to a specific paragraph in a Word document."""
        return footnote_tools.add_endnote_to_document(filename, paragraph_index, endnote_text)
    
    @mcp.tool()
    def customize_footnote_style(filename: str, numbering_format: str = "1, 2, 3",
                                start_number: int = 1, font_name: str = None,
                                font_size: int = None):
        """Customize footnote numbering and formatting in a Word document."""
        return footnote_tools.customize_footnote_style(
            filename, numbering_format, start_number, font_name, font_size
        )
    
    # Extended document tools
    @mcp.tool()
    def get_paragraph_text_from_document(filename: str, paragraph_index: int):
        """Get text from a specific paragraph in a Word document."""
        return extended_document_tools.get_paragraph_text_from_document(filename, paragraph_index)
    
    @mcp.tool()
    def find_text_in_document(filename: str, text_to_find: str, match_case: bool = True,
                             whole_word: bool = False):
        """Find occurrences of specific text in a Word document."""
        return extended_document_tools.find_text_in_document(
            filename, text_to_find, match_case, whole_word
        )
    
    @mcp.tool()
    def convert_to_pdf(filename: str, output_filename: str = None):
        """Convert a Word document to PDF format."""
        return extended_document_tools.convert_to_pdf(filename, output_filename)

    # 注册 process_file 工具（已替换 process_and_upload_file）
    @mcp.tool()
    def process_file(filename: Optional[str] = None, file_url: Optional[str] = None, process_type: str = "extract"):
        """
        支持本地文件或URL文件输入，解析后返回全部内容信息（文本、表格、图表、元信息等），不上传文件。
        process_type: 目前仅支持'extract'，后续可扩展。
        """
        import os
        import json
        from word_document_server.utils.file_utils import download_file_from_url
        from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure
        from docx import Document
        LOCAL_TMP_DIR = "tmp_files"
        # 1. 下载或定位本地文件
        if file_url is not None:
            local_path = download_file_from_url(file_url, LOCAL_TMP_DIR)
        elif filename is not None:
            local_path = filename
        else:
            return {"error": "必须提供filename或file_url"}
        if not os.path.exists(local_path):
            return {"error": f"文件不存在: {local_path}"}
        # 2. 解析文档内容
        try:
            doc = Document(local_path)
            # 文本内容
            text = "\n".join([p.text for p in doc.paragraphs])
            # 表格内容
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
            # 图表信息（简单提取嵌入图片）
            images = []
            for rel in doc.part.rels.values():
                if hasattr(rel, 'target_ref') and rel.target_ref and "image" in rel.target_ref:
                    images.append(rel.target_ref)
            # 元信息
            meta = get_document_properties(str(local_path))
            # 结构化大纲
            outline = get_document_structure(str(local_path))
            return {
                "text": text,
                "tables": tables,
                "images": images,
                "meta": meta,
                "outline": outline
            }
        except Exception as e:
            return {"error": f"解析文档失败: {str(e)}"}

    # 新增 create_document_and_upload 工具
    @mcp.tool()
    def create_document_and_upload(filename: str, title: str = None, author: str = None):
        """
        创建 Word 文档并上传到服务器，返回公网下载链接和服务器路径。
        """
        import os
        from word_document_server.tools import document_tools
        from word_document_server.utils.file_utils import upload_file_to_server
        # 1. 创建文档
        result = document_tools.create_document(filename, title, author)
        # 2. 检查创建是否成功
        if not (isinstance(result, str) and "created successfully" in result):
            return {"error": result}
        # 3. 上传到服务器
        REMOTE_DIR = "/root/files"
        SERVER = "8.156.74.79"
        USERNAME = "root"
        PASSWORD = "zfsZBC123."
        local_path = filename
        remote_path = os.path.join(REMOTE_DIR, os.path.basename(local_path))
        upload_result = upload_file_to_server(local_path, remote_path, SERVER, USERNAME, PASSWORD)
        public_url = f"http://8.156.74.79:8001/{os.path.basename(local_path)}"
        return {
            "message": result,
            "public_url": public_url,
            "remote_path": remote_path,
            "upload_result": upload_result
        }

    # 新增自动生成并上传Word文档的API（优化版本）
    @mcp.tool()
    async def auto_generate_and_upload_word(
        filename: str,
        content: dict
    ):
        """
        批量生成Word文档并上传到服务器 - 性能优化版本
        
        这个工具是接收JSON文件生成Word文档的核心工具，解决了传统方法中每次添加内容都要重新打开、修改、保存文档的性能问题。
        
        核心优化：
        1. 内存操作：文档在内存中处理，避免频繁的磁盘I/O操作
        2. 批量保存：所有内容添加完成后才保存一次，大幅提升性能
        3. 统计信息：提供详细的处理统计和错误信息
        4. 自动上传：生成完成后自动上传到服务器并返回公网链接
        
        大数据处理能力：
        - 支持处理包含数万条记录的JSON数据
        - 可处理包含数百个表格的大型数据集
        - 支持处理包含大量图片的复杂文档
        - 内存优化设计，避免大数据量处理时的性能瓶颈
        - 批量处理机制，大幅提升大数据量处理效率
        
        性能对比：
        传统方法：添加标题→保存→添加段落→保存→添加表格→保存...（N次I/O）
        优化方法：添加标题→添加段落→添加表格→...→保存→上传（1次I/O）
        
        大数据量处理优势：
        - 可处理包含数万条记录的JSON数据，性能提升5-10倍
        - 支持数百个表格的批量处理，避免内存溢出
        - 大量图片处理优化，减少磁盘I/O压力
        - 智能内存管理，自动清理资源
        
        Args:
            filename: 目标Word文件名（如 "report.docx"）
            content: 结构化内容字典，包含以下字段：
                - title: 文档标题（可选）
                - author: 文档作者（可选）
                - headings: 标题列表，每个元素包含text和level
                - paragraphs: 段落文本列表
                - tables: 表格数据列表，每个元素包含data字段
                - images: 图片列表，每个元素包含path和width
                - page_breaks: 分页符位置列表（可选）
        
        Returns:
            Dict[str, Any]: 包含以下字段的结果字典：
                - message: 处理结果描述
                - public_url: 公网下载链接
                - remote_path: 服务器上的文件路径
                - upload_result: 上传操作结果
                - stats: 文档生成统计信息（标题、段落、表格、图片数量等）
                - results: 详细的生成操作结果
                - error: 如果出错，包含错误信息
        
        使用示例：
        content = {
            "title": "项目报告",
            "author": "张三",
            "headings": [{"text": "第一章", "level": 1}],
            "paragraphs": ["这是第一段内容"],
            "tables": [{"data": [["列1", "列2"], ["数据1", "数据2"]]}]
        }
        result = await auto_generate_and_upload_word("report.docx", content)
        print(f"文档已上传，下载链接: {result['public_url']}")
        """
        return await batch_content_tools.batch_generate_and_upload_word(filename, content)

    # 新增批量生成Word文档API（不上传）
    @mcp.tool()
    async def batch_generate_word_document(
        filename: str,
        content: dict,
        save_after_batch: bool = True
    ):
        """
        批量生成Word文档 - 本地版本（不上传）
        
        这个工具专门用于批量生成Word文档，使用内存操作优化性能，但不包含上传功能。
        适用于只需要本地文档生成，不需要上传到服务器的场景。
        
        核心优化：
        1. 内存操作：文档在内存中处理，避免频繁的磁盘I/O操作
        2. 批量保存：所有内容添加完成后才保存一次，大幅提升性能
        3. 统计信息：提供详细的处理统计和错误信息
        4. 灵活控制：可以选择是否保存文档
        
        大数据处理能力：
        - 支持处理包含数万条记录的JSON数据
        - 可处理包含数百个表格的大型数据集
        - 支持处理包含大量图片的复杂文档
        - 内存优化设计，避免大数据量处理时的性能瓶颈
        - 批量处理机制，大幅提升大数据量处理效率
        
        性能优势：
        - 减少磁盘I/O操作次数：从N次（每次添加内容都保存）减少到1次（最后统一保存）
        - 提高处理速度：内存操作比磁盘操作快几个数量级
        - 降低系统负载：减少文件锁定和磁盘访问冲突
        - 支持大数据量处理：可高效处理包含数万条记录、数百个表格的大型JSON数据
        
        Args:
            filename: 目标Word文件名（如 "report.docx"）
            content: 结构化内容字典，包含以下字段：
                - title: 文档标题（可选）
                - author: 文档作者（可选）
                - headings: 标题列表，每个元素包含text和level
                - paragraphs: 段落文本列表
                - tables: 表格数据列表，每个元素包含data字段
                - images: 图片列表，每个元素包含path和width
                - page_breaks: 分页符位置列表（可选）
            save_after_batch: 是否在批量处理后保存文档（默认True）
        
        Returns:
            Dict[str, Any]: 包含以下字段的结果字典：
                - message: 处理结果描述
                - filename: 文件名
                - stats: 统计信息（标题、段落、表格、图片、分页符数量，错误列表）
                - results: 详细的操作结果列表
                - saved: 是否已保存
                - error: 如果出错，包含错误信息
        
        使用示例：
        content = {
            "title": "项目报告",
            "author": "张三",
            "headings": [{"text": "第一章", "level": 1}],
            "paragraphs": ["这是第一段内容"],
            "tables": [{"data": [["列1", "列2"], ["数据1", "数据2"]]}]
        }
        result = await batch_generate_word_document("report.docx", content)
        print(f"文档生成完成: {result['message']}")
        """
        return await batch_content_tools.batch_generate_word_document(filename, content, save_after_batch)


def run_server():
    """Run the Word Document MCP Server with configurable transport."""
    # Get transport configuration
    config = get_transport_config()
    
    # Setup logging
    # setup_logging(config['debug'])
    
    # Register all tools
    register_tools()
    
    # Print startup information
    transport_type = config['transport']
    print(f"Starting Word Document MCP Server with {transport_type} transport...")
    
    # if config['debug']:
    #     print(f"Configuration: {config}")
    
    try:
        if transport_type == 'stdio':
            # Run with stdio transport (default, backward compatible)
            print("Server running on stdio transport")
            mcp.run(transport='stdio')
            
        elif transport_type == 'streamable-http':
            # Run with streamable HTTP transport
            print(f"Server running on streamable-http transport at http://{config['host']}:{config['port']}{config['path']}")
            mcp.run(
                transport='streamable-http',
                host=config['host'],
                port=config['port'],
                path=config['path']
            )
            
        elif transport_type == 'sse':
            # Run with SSE transport
            print(f"Server running on SSE transport at http://{config['host']}:{config['port']}{config['sse_path']}")
            mcp.run(
                transport='sse',
                host=config['host'],
                port=config['port'],
                path=config['sse_path']
            )
            
    except KeyboardInterrupt:
        print("\nShutting down server...")
    except Exception as e:
        print(f"Error starting server: {e}")
        if config['debug']:
            import traceback
            traceback.print_exc()
        sys.exit(1)
    
    return mcp


def main():
    """Main entry point for the server."""
    run_server()


if __name__ == "__main__":
    main()